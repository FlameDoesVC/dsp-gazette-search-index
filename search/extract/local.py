"""Local extraction: the cheap rungs of the ladder. Spec 5.6.

Routing decision is chars-per-page, never file size. Measured on the corpus: a
58 KB single page was scanned while a 1.1 MB 28-page document had a complete
text layer, so any size heuristic is wrong.
"""

from __future__ import annotations

import io
import logging
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass

from django.conf import settings

logger = logging.getLogger(__name__)

TEXT_CAP = 20_000
_WS = re.compile(r"[ \t]+")
_BLANKS = re.compile(r"\n{3,}")


@dataclass(slots=True)
class ExtractionResult:
    text: str = ""
    page_count: int | None = None
    chars_per_page: int | None = None
    method: str = "none"
    status: str = "failed"
    error: str = ""
    transcribed: bool = False


def _clean(text: str) -> str:
    text = _WS.sub(" ", text or "")
    return _BLANKS.sub("\n\n", text).strip()[:TEXT_CAP]


def have_poppler() -> bool:
    return bool(shutil.which("pdftotext") and shutil.which("pdfinfo"))


def extract_docx(content: bytes) -> ExtractionResult:
    """python-docx, pure Python, flat memory. 73 files in the corpus and often
    the detail sheet -- the best value on the whole ladder."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = _clean("\n".join(parts))
    except Exception as exc:
        return ExtractionResult(
            method="docx", status="failed", error=str(exc)[:500]
        )
    if not text:
        return ExtractionResult(method="docx", status="failed", error="empty")
    return ExtractionResult(text=text, method="docx", status="ok")


def _page_count(path: str) -> int | None:
    try:
        out = subprocess.run(
            ["pdfinfo", path], capture_output=True, timeout=60, check=False
        ).stdout.decode("utf-8", "replace")
    except Exception:
        return None
    for line in out.splitlines():
        if line.startswith("Pages:"):
            try:
                return int(line.split()[1])
            except (IndexError, ValueError):
                return None
    return None


def extract_pdf_text_layer(content: bytes) -> ExtractionResult:
    """`pdftotext -layout`: a subprocess with a flat memory profile.

    `-layout` preserves column structure, which matters because the content we
    most want out of these files is a salary table.
    """
    if not have_poppler():
        return ExtractionResult(
            method="pdftotext", status="failed", error="poppler not installed"
        )
    with tempfile.NamedTemporaryFile(suffix=".pdf") as handle:
        handle.write(content)
        handle.flush()
        pages = _page_count(handle.name)
        try:
            proc = subprocess.run(
                ["pdftotext", "-layout", "-q", handle.name, "-"],
                capture_output=True,
                timeout=120,
                check=False,
            )
            raw = proc.stdout.decode("utf-8", "replace")
        except Exception as exc:
            return ExtractionResult(
                page_count=pages, method="pdftotext",
                status="failed", error=str(exc)[:500],
            )

    text = _clean(raw)
    dense = len(re.sub(r"\s", "", text))
    per_page = int(dense / pages) if pages else (dense if text else 0)
    return ExtractionResult(
        text=text,
        page_count=pages,
        chars_per_page=per_page,
        method="pdftotext",
        status="ok",
    )


def needs_transcription(result: ExtractionResult) -> bool:
    """True when the PDF has no usable text layer and must go to a vision
    model. Routing is on extracted density alone."""
    if result.status != "ok":
        return True
    threshold = getattr(settings, "SCANNED_CHARS_PER_PAGE", 200)
    return (result.chars_per_page or 0) < threshold
